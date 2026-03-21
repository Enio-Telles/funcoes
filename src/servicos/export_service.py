from __future__ import annotations

import html
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

import polars as pl
from docx import Document

from src.config import MAX_DOCX_ROWS
from src.utilitarios.text import display_cell

class ExportService:
    @staticmethod
    def _iter_rows(df: pl.DataFrame):
        """
        Otimização: Substituído o loop manual do python na aplicação do display_cell
        por operações nativas do Polars, resultando em conversão muito mais eficiente.
        """
        cols = []
        for c in df.columns:
            dtype = df[c].dtype
            if isinstance(dtype, (pl.List, pl.Struct, pl.Object)):
                cols.append(pl.col(c).map_elements(display_cell, return_dtype=pl.String).fill_null("").alias(c))
            elif dtype == pl.Boolean:
                cols.append(
                    pl.when(pl.col(c))
                    .then(pl.lit("true"))
                    .when(~pl.col(c))
                    .then(pl.lit("false"))
                    .otherwise(pl.lit(""))
                    .alias(c)
                )
            else:
                cols.append(pl.col(c).cast(pl.String).fill_null("").alias(c))

        df_str = df.select(cols)
        for row in df_str.iter_rows():
            yield row

    def export_excel(self, target: Path, df: pl.DataFrame, sheet_name: str = "Dados") -> Path:
        """
        Otimização: Substituído o loop manual do openpyxl pelo motor nativo do Polars.
        Utiliza write_excel com autofit, muito mais rápido para grandes volumes.
        """
        target.parent.mkdir(parents=True, exist_ok=True)
        # Polars usa o xlsxwriter ou calamine por baixo (se disponível)
        # Nota: worksheet name no Excel tem limite de 31 caracteres
        df.write_excel(
            target, 
            worksheet=sheet_name[:31], 
            table_style="Table Style Medium 9",
            autofit=True
        )
        return target

    def build_html_report(
        self,
        title: str,
        cnpj: str,
        table_name: str,
        df: pl.DataFrame,
        filters_text: str,
        visible_columns: list[str],
    ) -> str:
        headers = "".join(f"<th>{html.escape(col)}</th>" for col in df.columns)
        body_rows = []
        for row in self._iter_rows(df):
            cells = "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row)
            body_rows.append(f"<tr>{cells}</tr>")
        body = "\n".join(body_rows)
        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        return f"""<!DOCTYPE html>
<html lang=\"pt-BR\">
<head>
<meta charset=\"utf-8\">
<title>{html.escape(title)}</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 24px; color: #1f2937; }}
h1, h2 {{ margin-bottom: 8px; }}
.meta {{ background: #f5f7fb; border: 1px solid #dbe2ea; padding: 12px; border-radius: 8px; margin-bottom: 18px; }}
table {{ border-collapse: collapse; width: 100%; font-size: 12px; }}
th, td {{ border: 1px solid #d1d5db; padding: 6px 8px; text-align: left; vertical-align: top; }}
th {{ background: #eef2f7; position: sticky; top: 0; }}
</style>
</head>
<body>
<h1>{html.escape(title)}</h1>
<div class=\"meta\">
<p><strong>CNPJ:</strong> {html.escape(cnpj)}</p>
<p><strong>Tabela:</strong> {html.escape(table_name)}</p>
<p><strong>Gerado em:</strong> {generated_at}</p>
<p><strong>Filtros:</strong> {html.escape(filters_text or 'Sem filtros')}</p>
<p><strong>Colunas visíveis:</strong> {html.escape(', '.join(visible_columns) if visible_columns else 'Todas')}</p>
<p><strong>Linhas no relatório:</strong> {df.height}</p>
</div>
<h2>Dados</h2>
<table>
<thead><tr>{headers}</tr></thead>
<tbody>
{body}
</tbody>
</table>
</body>
</html>"""

    def export_txt_with_html(self, target: Path, html_report: str) -> Path:
        """
        Segurança: Gravação atômica via arquivo temporário.
        Evita corrupção de arquivos em caso de falha durante a escrita.
        """
        target.parent.mkdir(parents=True, exist_ok=True)
        
        with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', delete=False, suffix='.tmp') as tf:
            tf.write(html_report)
            temp_path = Path(tf.name)
            
        try:
            shutil.move(str(temp_path), str(target))
        except Exception as e:
            if temp_path.exists():
                temp_path.unlink()
            raise e
            
        return target

    def export_docx(
        self,
        target: Path,
        title: str,
        cnpj: str,
        table_name: str,
        df: pl.DataFrame,
        filters_text: str,
        visible_columns: list[str],
    ) -> Path:
        target.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"CNPJ: {cnpj}")
        doc.add_paragraph(f"Tabela: {table_name}")
        doc.add_paragraph(f"Filtros: {filters_text or 'Sem filtros'}")
        doc.add_paragraph(f"Colunas visíveis: {', '.join(visible_columns) if visible_columns else 'Todas'}")
        doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph(f"Linhas incluídas: {df.height}")

        rows_to_write = min(df.height, MAX_DOCX_ROWS)
        if df.height > MAX_DOCX_ROWS:
            doc.add_paragraph(
                f"Observação: por desempenho, o relatório Word inclui as primeiras {MAX_DOCX_ROWS} linhas do recorte selecionado."
            )
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for idx, col in enumerate(df.columns):
            header_cells[idx].text = str(col)
            
        # Otimização: iter_rows() sem named=True
        for row in self._iter_rows(df.head(rows_to_write)):
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = value
                
        doc.save(target)
        return target
